import io
import docx
from typing import Union
from backend.ingestion.models.upload_dtos import ParsedDocumentDTO

def parse_docx(file_data: Union[bytes, str], file_name: str) -> ParsedDocumentDTO:
    if isinstance(file_data, bytes):
        doc = docx.Document(io.BytesIO(file_data))
    else:
        doc = docx.Document(file_data)
        
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return ParsedDocumentDTO(file_name=file_name, file_type="docx", content=text)
